import sys
import argparse
from datetime import datetime
from pathlib import Path
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
import matplotlib as mpl
from matplotlib.patches import Patch

import pandas as pd
from urllib.parse import urlparse
from docx import Document

from utils import load_config, init_logger
from gsc_fetcher import authenticate, list_properties, fetch_performance
from analyzer import init_analyzer, compute_summary, segment_dataframe, detect_low_hanging, compute_mom, detect_anomalies
from visualizer import init_visualizer, plot_pie, plot_multi_line

logger = None

def select_property(props):
    print("Available GSC Properties:")
    for i, p in enumerate(props, 1):
        print(f"  {i}. {p}")
    choice = input("Select property number: ")
    try:
        return props[int(choice) - 1]
    except:
        logger.error("Invalid selection")
        sys.exit(1)

def get_folder(path):
    parsed = urlparse(path)
    parts = [seg for seg in parsed.path.split('/') if seg]
    return parts[0] if parts else '/'

def build_report(cfg, service, logger, site_url):
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_dir = Path(cfg['output']['excel_path']).parent
    out_dir.mkdir(parents=True, exist_ok=True)

    excel_path = out_dir / f"{Path(cfg['output']['excel_path']).stem}_{ts}.xlsx"
    md_path    = out_dir / f"{Path(cfg['output']['markdown_path']).stem}_{ts}.md"
    docx_path  = out_dir / f"{Path(cfg['output']['docx_path']).stem}_{ts}.docx"

    # Auto-detect end_date
    if not cfg['dates']['end_date']:
        df_dates = fetch_performance(
            service, logger, site_url,
            cfg['dates']['start_date'], datetime.utcnow().strftime("%Y-%m-%d"),
            ['date']
        )
        cfg['dates']['end_date'] = (
            df_dates['date'].max()
            if not df_dates.empty else cfg['dates']['start_date']
        )
        logger.info(f"Detected end_date: {cfg['dates']['end_date']}")

    # Country filter
    base_filters = []
    country = cfg['filters'].get('country','')
    if country:
        base_filters.append({
            'dimension':'country','operator':'equals','expression':country
        })

    # 1) Fetch FULL dataset
    df_full = fetch_performance(
        service, logger, site_url,
        cfg['dates']['start_date'], cfg['dates']['end_date'],
        ['page','query'], filters=base_filters
    )
    for col in ['clicks','impressions','ctr','position']:
        df_full[col] = pd.to_numeric(df_full[col], errors='coerce').fillna(0)
    df_full['query'] = df_full['query'].fillna("").astype(str)

    # 2) Segment using your regex
    sum_o = compute_summary(df_full, 'Overall')
    branded, nonb, anon = segment_dataframe(df_full, cfg['branded']['regex'])
    sum_b  = compute_summary(branded,   'Branded')
    sum_nb = compute_summary(nonb,      'Non-Branded')
    sum_an = compute_summary(anon,      'Anonymous')

    # 3) Top Pages & Queries
    top_pages = (df_full.groupby('page')
                    .agg(clicks=('clicks','sum'),
                         impressions=('impressions','sum'),
                         ctr=('ctr','mean'),
                         position=('position','mean'))
                    .reset_index()
                    .nlargest(20,'clicks'))
    top_queries = (df_full.groupby('query')
                      .agg(clicks=('clicks','sum'),
                           impressions=('impressions','sum'),
                           ctr=('ctr','mean'),
                           position=('position','mean'))
                      .reset_index()
                      .nlargest(20,'clicks'))

    # 4) Folder analysis
    folder_summaries = {}
    for label, df_seg in [('Overall',df_full),('Branded',branded),('Non-Branded',nonb)]:
        try:
            tmp = df_seg.copy()
            tmp['folder'] = tmp['page'].apply(get_folder)
            fs = (tmp.groupby('folder')
                     .agg(clicks=('clicks','sum'),
                          impressions=('impressions','sum'))
                     .reset_index()
                     .nlargest(10,'clicks'))
        except Exception as e:
            logger.warning(f"Folder analysis failed for {label}: {e}")
            fs = pd.DataFrame(columns=['folder','clicks','impressions'])
        folder_summaries[label] = fs
    
    records = []
    for page, grp in df_full.groupby('page'):
        segs = [s for s in urlparse(page).path.split('/') if s]
        # metrics for this page
        clicks = grp['clicks'].sum()
        impr   = grp['impressions'].sum()
        ctr    = clicks / impr if impr else 0
        pos    = (grp['position'] * grp['impressions']).sum() / impr if impr else 0
        # generate one record per folder‐level
        for lvl in range(1, len(segs)+1):
            folder = "/".join(segs[:lvl])
            records.append({
                'folder':        folder,
                'page':          page,
                'clicks':        clicks,
                'impressions':   impr,
                'ctr':           ctr,
                'avg_position':  pos
            })

    df_folders = pd.DataFrame(records)

    # Summary: aggregate per folder
    df_folder_summary = (
        df_folders
        .groupby('folder')
        .agg(
            url_count    = ('page', 'nunique'),
            clicks       = ('clicks','sum'),
            impressions  = ('impressions','sum'),
            avg_ctr      = ('ctr','mean'),
            avg_position = ('avg_position','mean')
        )
        .reset_index()
        .sort_values('clicks', ascending=False)
    )

    # Detail: list URLs under each folder
    df_folder_urls = (
        df_folders
        .groupby('folder')['page']
        .unique()
        .reset_index()
        .rename(columns={'page':'urls'})
    )
    df_folder_urls['url_count'] = df_folder_urls['urls'].apply(len)


    # 5) MoM
    df_dq = fetch_performance(
        service, logger, site_url,
        cfg['dates']['start_date'], cfg['dates']['end_date'],
        ['date','query'], filters=base_filters
    )
    mom_o  = compute_mom(df_dq)
    mom_b  = compute_mom(df_dq[df_dq['query'].str.contains(cfg['branded']['regex'], regex=True)])
    mom_nb = compute_mom(df_dq[~df_dq['query'].str.contains(cfg['branded']['regex'], regex=True)])
    for df in (mom_o,mom_b,mom_nb):
        df.insert(0, 'month_label', df.pop('month_label'))
    
    daily = (
    df_dq
    .groupby('date')
    .agg(
        clicks=('clicks','sum'),
        impressions=('impressions','sum'),
        ctr=('ctr','mean'),
        position=('position','mean')
    )
    .reset_index()
    )
    
    # After df_dq is fetched and numeric-cast:
    anoms_clicks = detect_anomalies(daily, date_col='date', metric='clicks', window=7, z_thresh=2.5)
    anoms_impr   = detect_anomalies(daily, date_col='date', metric='impressions', window=7, z_thresh=2.5)

    # 6) Export everything, including raw tabs
    with pd.ExcelWriter(excel_path, engine='openpyxl') as w:
        df_full.to_excel(w, 'RawFull', index=False)
        branded.to_excel(w, 'RawBranded', index=False)
        nonb.to_excel(w, 'RawNonBranded', index=False)

        pd.DataFrame([sum_o,sum_b,sum_nb,sum_an]).to_excel(w, 'Summary', index=False)

        avg_df = pd.DataFrame([
            {'segment':'Overall',    **mom_o[['clicks','impressions','ctr','avg_position']].mean().to_dict()},
            {'segment':'Branded',    **mom_b[['clicks','impressions','ctr','avg_position']].mean().to_dict()},
            {'segment':'Non-Branded',**mom_nb[['clicks','impressions','ctr','avg_position']].mean().to_dict()}
        ]).round({'ctr':4,'avg_position':2})
        avg_df.to_excel(w, 'MonthlyAverages', index=False)

        top_pages.to_excel(w, 'TopPages', index=False)
        top_queries.to_excel(w, 'TopQueries', index=False)

        for lbl, df_f in folder_summaries.items():
            df_f.to_excel(w, f'Folders_{lbl}', index=False)

        mom_o.to_excel(w, 'MoM_Overall', index=False)
        mom_b.to_excel(w, 'MoM_Branded', index=False)
        mom_nb.to_excel(w, 'MoM_NonBranded', index=False)
        anoms_clicks.to_excel(w, 'Anomalies_Clicks', index=False)
        anoms_impr.to_excel(w,   'Anomalies_Impressions', index=False)

        # 6) Low-hanging opportunities (derive directly from df_full)
        low_hanging = detect_low_hanging(df_full, **cfg['thresholds']['low_hanging'])
        low_hanging.to_excel(w, "LowHanging", index=False)

        df_folder_summary.to_excel(w, 'Folders_Multi', index=False)

        df_folder_urls['urls'] = df_folder_urls['urls'].apply(lambda lst: "\n".join(lst))
        df_folder_urls.to_excel(w, 'Folder_URLs', index=False)


    #     detect_low_hanging(df_full, **cfg['thresholds']['low_hanging']).to_excel(w, 'LowHanging', index=False)
    # logger.info(f"Excel saved: {excel_path}")

    # 7) Charts
    segments = {'Overall':mom_o,'Branded':mom_b,'Non-Branded':mom_nb}
    if cfg['visualization']['line_charts']:
        for metric in ['clicks','impressions','ctr','avg_position']:
            plot_multi_line(segments,'month_label',metric,
                            f"MoM_{metric.capitalize()}_By_Segment",str(out_dir))
    if cfg['visualization']['pie_charts']:
        filters_map = [
            ('Overall', None),
            ('Branded',    {'dimension':'query','operator':'contains','expression':cfg['branded']['regex']}),
            ('Non-Branded',{'dimension':'query','operator':'notContains','expression':cfg['branded']['regex']})
        ]
        for label, fp in filters_map:
            seg_f = base_filters.copy()
            if fp: seg_f.append(fp)
            df_dev = fetch_performance(service,logger,site_url,
                                       cfg['dates']['start_date'],cfg['dates']['end_date'],
                                       ['device'],filters=seg_f)
            if not df_dev.empty:
                plot_pie(df_dev,'device','clicks',f"Device_{label}",str(out_dir))
            else:
                logger.warning(f"No device data for '{label}'")
        seg_counts = pd.DataFrame([
            {'segment':'Branded','clicks':sum_b['clicks']},
            {'segment':'Non-Branded','clicks':sum_nb['clicks']}
           # {'segment':'Overall','clicks':sum_o['clicks']}
        ])
        plot_pie(seg_counts,'segment','clicks',"Segment_Clicks",str(out_dir))
    
    fig, ax = plt.subplots(figsize=(12, 6))
    # Color bars: red if anomaly, gray otherwise
    anom_dates = set(anoms_clicks['date'])
    colors = ['red' if d in anom_dates else 'gray' for d in daily['date']]

    ax.bar(daily['date'], daily['clicks'], color=colors, width=0.8)

    # One tick per week for clarity
    ax.xaxis.set_major_locator(mdates.WeekdayLocator(byweekday=mdates.MO))
    ax.xaxis.set_major_formatter(mdates.DateFormatter('%b %d'))

    ax.set_title('Daily Clicks with Anomaly Flags')
    ax.set_xlabel('Date')
    ax.set_ylabel('Clicks')
    plt.xticks(rotation=45, ha='right')

    # Legend
    handles = [Patch(color='gray', label='Normal'),
            Patch(color='red',  label='Anomaly')]
    ax.legend(handles=handles, loc='upper left')

    # Save
    anom_chart_path = out_dir / f"anomalies_clicks_{ts}.png"
    fig.savefig(str(anom_chart_path), bbox_inches='tight')
    plt.close(fig)
    logger.info(f"Saved anomaly bar chart: {anom_chart_path}")


    # 8) Actionable Insights
    drop = mom_o[mom_o['pct_clicks']<0]
    worst = drop.loc[drop['pct_clicks'].idxmin()] if not drop.empty else None
    low_count = len(detect_low_hanging(df_full,**cfg['thresholds']['low_hanging']))

    if cfg['output']['formats']['markdown']:
        with open(md_path,'w') as md:
            md.write(f"# GSC Audit Report for {site_url}\n\n")
            for s in [sum_o,sum_b,sum_nb]:
                md.write(f"- **{s['segment']}**: Clicks={s['clicks']}, Impr={s['impressions']}, "
                         f"CTR={s['ctr']:.2%}, AvgPos={s['avg_position']:.2f}\n")
            md.write("\n## Actionable Insights\n")
            if worst is not None:
                md.write(f"- Largest MoM click drop: {worst['month_label']} ({worst['pct_clicks']:.2%})\n")
            else:
                md.write("- No negative MoM click changes detected.\n")
            md.write(f"- {low_count} low-hanging opportunities identified.\n")
        logger.info(f"Markdown saved: {md_path}")

    if cfg['output']['formats']['docx']:
        try:
            doc = Document()
            doc.add_heading(f"GSC Audit Report for {site_url}", 0)
            doc.add_heading("Performance Summary", level=1)
            for s in [sum_o,sum_b,sum_nb]:
                doc.add_paragraph(f"{s['segment']}: Clicks={s['clicks']}, Impr={s['impressions']}, "
                                  f"CTR={s['ctr']:.2%}, AvgPos={s['avg_position']:.2f}")
            doc.add_heading("Actionable Insights", level=1)
            if worst is not None:
                doc.add_paragraph(f"Largest MoM click drop: {worst['month_label']} ({worst['pct_clicks']:.2%})")
            else:
                doc.add_paragraph("No negative MoM click changes detected.")
            doc.add_paragraph(f"Low-hanging opportunities: {low_count} rows.")
            doc.save(docx_path)
            logger.info(f"Word report saved: {docx_path}")
        except Exception as e:
            logger.warning(f"Word export failed: {e}")

def main():
    parser = argparse.ArgumentParser(description="GSC Audit Automation Tool")
    parser.add_argument('--config', default='config.yaml', help='Path to config YAML')
    parser.add_argument('--property', help='Site URL to audit')
    args = parser.parse_args()

    cfg = load_config(args.config)
    global logger
    logger = init_logger(cfg['logging']['file'], cfg['logging']['level'])

    service, logger = authenticate()
    if cfg['interactive'] and not args.property:
        props = list_properties(service, logger)
        site_url = select_property(props)
    else:
        site_url = args.property
    if not site_url:
        logger.error("No site property provided.")
        sys.exit(1)

    init_analyzer(cfg)
    init_visualizer(cfg)
    build_report(cfg, service, logger, site_url)

if __name__ == '__main__':
    main()
